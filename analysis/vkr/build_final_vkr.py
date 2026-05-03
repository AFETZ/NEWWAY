#!/usr/bin/env python3
"""Build final VKR DOCX from the expanded source package.

The script keeps the user's working files intact by building a new master
document in output/doc/. It uses the existing VKR DOCX as the body source,
prepends newly assembled front matter, injects figures from analysis/vkr/figures,
and appends supporting appendices before the bibliography section.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/home/afetz/work/clean/NEWWAY")
BASE_DOCX = ROOT / "VKR_corrected_1_2_3_zak.docx"
FRONT_MATTER = ROOT / "analysis" / "vkr" / "VKR_front_matter.md"
APPENDICES = ROOT / "analysis" / "vkr" / "VKR_appendices.md"
FIG_DIR = ROOT / "analysis" / "vkr" / "figures"
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_DOCX = OUTPUT_DIR / "Fizulin_AV_VKR_final_gost_7_32.docx"
BIB_SOURCE = ROOT / "analysis" / "vkr" / "VKR_bibliography_verified.md"
TOC_STATS = ROOT / "output" / "pdf" / "Fizulin_AV_VKR_final_gost_7_32.stats.json"


HEADING_RE = re.compile(r"^(#+)\s+(.*)$")
TOC_RE = re.compile(r"^\[TOC\]$")
TABLE_RE = re.compile(r"^\|")
ORDERED_RE = re.compile(r"^\d+\.\s+")
CAPTION_RE = re.compile(r"^(Рисунок|Таблица|Приложение)\s+")

FIGURE_MAP = {
    "Рисунок 2.1 — Архитектура инструментальной среды совместного моделирования": FIG_DIR / "figure_2_1_architecture.png",
    "Рисунок 3.1 — Динамика скоростей и полос движения в безопасном режиме сценария объезда остановившегося лидирующего транспортного средства": FIG_DIR / "figure_3_1_speed_lane_safe.png",
    "Рисунок 3.2 — Накопительный PRR предупредительных CAM-сообщений в безопасном режиме": FIG_DIR / "figure_3_2_prr_cumulative_safe.png",
    "Рисунок 3.3 — Цепочка “эквивалентная мощность — PRR — маневр” в безопасном режиме": FIG_DIR / "figure_3_3_dbm_prr_chain_safe.png",
    "Рисунок 3.4 — Динамика скоростей и полос движения в аварийном режиме сценария объезда препятствия": FIG_DIR / "figure_3_4_speed_lane_crash.png",
    "Рисунок 3.5 — Поток сетевых событий и решений в аварийном режиме": FIG_DIR / "figure_3_5_events_crash.png",
    "Рисунок 3.6 — Причинная шкала событий в аварийном режиме": FIG_DIR / "figure_3_6_event_chain_crash.png",
    "Рисунок 3.7 — Наглядная цепочка “качество канала — PRR — решение — исход” для сценария объезда препятствия": FIG_DIR / "figure_3_7_dbm_prr_chain_crash.png",
    "Рисунок 3.8 — XY-траектории транспортных средств на перекрестке в трех режимах": FIG_DIR / "figure_3_8_xy_trajectories.png",
    "Рисунок 3.9 — Скоростной профиль veh3 и моменты реакции в сценарии перекрестка": FIG_DIR / "figure_3_9_veh3_speed.png",
    "Рисунок 3.10 — Доставка CAM от veh2 к veh3 и накопительный PRR в режимах radar_bad_link, radar_only и radar_good_link": FIG_DIR / "figure_3_10_comm_delivery.png",
    "Рисунок 3.11 — Нулевая разница скоростей между режимами до первого вычисленного управляющего воздействия": FIG_DIR / "figure_3_11_pre_event_spread.png",
    "Рисунок 3.12 — Параметрическая развертка по veh3_equiv_tx_power_dbm: влияние качества канала на PRR, момент первой реакции и исход на перекрестке": FIG_DIR / "figure_3_12_dbm_sweep.png",
}


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def insert_paragraph_before(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if text:
        paragraph.add_run(text)
    if style:
        paragraph.style = style
    return paragraph


def insert_table_before(anchor: Paragraph, rows: list[list[str]]) -> None:
    if not rows:
        return
    scratch = Document()
    table = scratch.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = cleanup_inline(value)
    anchor._p.addprevious(deepcopy(table._tbl))


def cleanup_inline(text: str) -> str:
    cleaned = text.replace("**", "").replace("`", "")
    cleaned = cleaned.replace("->", "->")
    return cleaned.strip()


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    idx = start
    while idx < len(lines) and TABLE_RE.match(lines[idx].strip()):
        table_lines.append(lines[idx].strip())
        idx += 1
    rows: list[list[str]] = []
    for raw in table_lines:
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(set(cell) <= {"-"} for cell in rows[1]):
        rows.pop(1)
    return rows, idx


def flush_paragraph(anchor: Paragraph, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip())
    text = cleanup_inline(text)
    if text:
        para = insert_paragraph_before(anchor, text, style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    buffer.clear()


def add_page_break_before(anchor: Paragraph) -> None:
    page_break = insert_paragraph_before(anchor, style="Normal")
    page_break.add_run().add_break(WD_BREAK.PAGE)


def add_toc_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run = paragraph.add_run("Обновите поле оглавления в Word/LibreOffice")
    run.italic = True
    run._r.append(fld_char_end)


def load_toc_entries() -> list[dict[str, str | int]]:
    if not TOC_STATS.exists():
        return []
    payload = json.loads(TOC_STATS.read_text(encoding="utf-8"))
    entries = []
    for item in payload.get("headings", []):
        text = str(item.get("text", "")).strip()
        page = item.get("page")
        level = int(item.get("level", 1))
        if not text or text == "Содержание" or page is None or level > 3:
            continue
        entries.append({"text": text, "page": int(page), "level": level})
    return entries


def add_static_toc(anchor: Paragraph, entries: list[dict[str, str | int]]) -> None:
    if not entries:
        para = insert_paragraph_before(anchor, style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_toc_field(para)
        return

    right_tab = Inches(6.1)
    for entry in entries:
        para = insert_paragraph_before(anchor, style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = {1: Cm(0), 2: Cm(1.25), 3: Cm(2.5)}.get(int(entry["level"]), Cm(0))
        pf.tab_stops.add_tab_stop(right_tab, alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        run = para.add_run(f'{entry["text"]}\t{entry["page"]}')
        if int(entry["level"]) == 1:
            run.bold = True


def find_paragraph(doc: Document, predicate) -> Paragraph:
    for para in doc.paragraphs:
        if predicate(para):
            return para
    raise ValueError("Anchor paragraph not found")


def add_title_page(doc: Document, anchor: Paragraph) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)

    lines = [
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАТЕЛЬНОЕ",
        "УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ",
        "«НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ «ВЫСШАЯ ШКОЛА ЭКОНОМИКИ»",
        "МОСКОВСКИЙ ИНСТИТУТ ЭЛЕКТРОНИКИ И МАТЕМАТИКИ им. А.Н. ТИХОНОВА",
        "",
        "Физулин Андрей Вадимович, группа МИВ241",
        "",
        "ИССЛЕДОВАНИЕ ВЛИЯНИЯ ПОТЕРЬ СООБЩЕНИЙ 5G NR",
        "НА ПОВЕДЕНИЕ ПОДКЛЮЧЕННЫХ И БЕСПИЛОТНЫХ",
        "ТРАНСПОРТНЫХ СРЕДСТВ",
        "",
        "Выпускная квалификационная работа",
        "магистерская диссертация",
        "по направлению 11.04.02 «Инфокоммуникационные технологии и системы связи»",
        "образовательная программа «Интернет вещей и киберфизические системы»",
        "",
        "Студент: Физулин Андрей Вадимович",
        "Научный руководитель: доцент, к.т.н. И.А. Иванов",
        "Со-руководитель: преподаватель ДКИ МИЭМ НИУ ВШЭ, В.Г. Степанянц",
        "",
        "",
        "Москва, 2026",
    ]
    for line in lines:
        para = insert_paragraph_before(anchor, line, style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line.startswith("ИССЛЕДОВАНИЕ"):
            para.runs[0].bold = True
        if line == "Выпускная квалификационная работа":
            para.runs[0].bold = True
    add_page_break_before(anchor)


def insert_markdown_before(doc: Document, anchor: Paragraph, md_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    buffer: list[str] = []
    idx = 0

    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped:
            flush_paragraph(anchor, buffer)
            idx += 1
            continue

        if TOC_RE.match(stripped):
            flush_paragraph(anchor, buffer)
            add_static_toc(anchor, load_toc_entries())
            idx += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph(anchor, buffer)
            level = len(heading_match.group(1))
            text = cleanup_inline(heading_match.group(2))
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
            if style == "Heading 1":
                add_page_break_before(anchor)
            para = insert_paragraph_before(anchor, text, style=style)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if style == "Heading 1" and not text[0].isdigit() else WD_ALIGN_PARAGRAPH.LEFT
            idx += 1
            continue

        if TABLE_RE.match(stripped):
            flush_paragraph(anchor, buffer)
            rows, idx = parse_markdown_table(lines, idx)
            insert_table_before(anchor, rows)
            continue

        if stripped.startswith("- "):
            flush_paragraph(anchor, buffer)
            while idx < len(lines) and lines[idx].strip().startswith("- "):
                item = cleanup_inline(lines[idx].strip()[2:])
                para = insert_paragraph_before(anchor, item, style="List Paragraph")
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                idx += 1
            continue

        if ORDERED_RE.match(stripped):
            flush_paragraph(anchor, buffer)
            while idx < len(lines) and ORDERED_RE.match(lines[idx].strip()):
                item = cleanup_inline(lines[idx].strip())
                para = insert_paragraph_before(anchor, item, style="List Paragraph")
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                idx += 1
            continue

        buffer.append(stripped)
        idx += 1

    flush_paragraph(anchor, buffer)


def replace_bibliography(doc: Document, heading_para: Paragraph, source_md: Path) -> None:
    if not source_md.exists():
        return

    entries: list[str] = []
    for raw in source_md.read_text(encoding="utf-8").splitlines():
        stripped = raw.strip()
        if re.match(r"^\d+\.\s+", stripped):
            entries.append(re.sub(r"^\d+\.\s+", "", stripped))

    start = next(i for i, para in enumerate(doc.paragraphs) if para._p is heading_para._p) + 1
    bib_paras = [p for p in doc.paragraphs[start:] if p.text.strip()]
    for para, text in zip(bib_paras, entries):
        para.text = text


def insert_figure_after(paragraph: Paragraph, image_path: Path, width_cm: float = 16.0) -> None:
    if not image_path.exists():
        return
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    image_para = Paragraph(new_p, paragraph._parent)
    image_para.style = "Normal"
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Cm(width_cm))


def inject_figures(doc: Document) -> None:
    inserted: set[str] = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in FIGURE_MAP and text not in inserted:
            insert_figure_after(para, FIGURE_MAP[text])
            inserted.add(text)


def ensure_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)


def build() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(BASE_DOCX))
    ensure_base_styles(doc)

    first_heading = find_paragraph(doc, lambda p: p.text.strip().startswith("1 Теоретические основы"))
    bibliography_heading = find_paragraph(doc, lambda p: p.text.strip() == "Список использованных источников")

    add_title_page(doc, first_heading)
    insert_markdown_before(doc, first_heading, FRONT_MATTER)
    insert_markdown_before(doc, bibliography_heading, APPENDICES)
    replace_bibliography(doc, bibliography_heading, BIB_SOURCE)
    inject_figures(doc)
    set_update_fields_on_open(doc)
    doc.save(str(OUTPUT_DOCX))
    return OUTPUT_DOCX


if __name__ == "__main__":
    result = build()
    print(result)
