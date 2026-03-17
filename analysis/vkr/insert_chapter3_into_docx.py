#!/usr/bin/env python3
"""Insert the practical chapter draft into an existing DOCX report."""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph


ROOT = Path("/home/afetz/work/clean/NEWWAY")
SOURCE_DOCX = ROOT / "Fizulin_Practic_Otchet_2026_v1.docx"
CHAPTER_MD = ROOT / "analysis" / "vkr" / "VKR_chapter3_practical_part_draft.md"
OUTPUT_DOCX = ROOT / "Fizulin_Practic_Otchet_2026_v2_practical.docx"

CAPTION_RE = re.compile(r"^\*\*(Рисунок\s+\d+\.\d+\s+—\s+.+)\*\*$")
PLACEHOLDER_RE = re.compile(r"^\[Вставить\s+(рисунок\s+\d+\.\d+)\s+здесь\]$", re.IGNORECASE)
HEADING_RE = re.compile(r"^(#+)\s+(.*)$")
NUMBERED_RE = re.compile(r"^\d+\.\s+")


def cleanup_inline_markup(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text.strip()


def extract_figure_map(lines: list[str]) -> dict[str, Path]:
    mapping: dict[str, Path] = {}
    for line in lines:
        if not line.startswith("| Рисунок"):
            continue
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        if len(parts) != 3:
            continue
        figure_id = parts[0].lower()
        raw_path = parts[2].strip()
        if raw_path.startswith("`") and raw_path.endswith("`"):
            raw_path = raw_path[1:-1]
        mapping[figure_id] = Path(raw_path)
    return mapping


def insert_paragraph_before(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if text:
        paragraph.add_run(text)
    if style:
        paragraph.style = style
    return paragraph


def insert_table_before(anchor: Paragraph, rows: list[list[str]]) -> None:
    if not rows:
        return
    scratch = Document()
    table = scratch.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = cleanup_inline_markup(value)
    anchor._p.addprevious(deepcopy(table._tbl))


def find_paragraph_by_text(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"Paragraph '{text}' not found")


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def flush_paragraph_buffer(anchor: Paragraph, buffer: list[str]) -> None:
    if not buffer:
        return
    text = cleanup_inline_markup(" ".join(part.strip() for part in buffer if part.strip()))
    if text:
        para = insert_paragraph_before(anchor, text, style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    buffer.clear()


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1

    rows: list[list[str]] = []
    for raw in table_lines:
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        rows.append(cells)

    if len(rows) >= 2 and all(set(cell) <= {"-"} for cell in rows[1]):
        rows.pop(1)
    return rows, idx


def build_docx() -> Path:
    doc = Document(str(SOURCE_DOCX))
    anchor = find_paragraph_by_text(doc, "Заключение")

    lines = CHAPTER_MD.read_text().splitlines()
    figure_map = extract_figure_map(lines)

    # Page break before the new chapter.
    page_break = insert_paragraph_before(anchor, style="Normal")
    page_break.add_run().add_break(WD_BREAK.PAGE)

    paragraph_buffer: list[str] = []
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if stripped == "":
            flush_paragraph_buffer(anchor, paragraph_buffer)
            idx += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph_buffer(anchor, paragraph_buffer)
            level = len(heading_match.group(1))
            text = cleanup_inline_markup(heading_match.group(2))
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
            para = insert_paragraph_before(anchor, text, style=style)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            idx += 1
            continue

        placeholder_match = PLACEHOLDER_RE.match(stripped)
        if placeholder_match:
            flush_paragraph_buffer(anchor, paragraph_buffer)
            figure_id = placeholder_match.group(1).lower()
            image_path = figure_map.get(figure_id)
            image_para = insert_paragraph_before(anchor, style="Normal")
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if image_path and image_path.exists():
                image_para.add_run().add_picture(str(image_path), width=Cm(16))
            else:
                image_para.add_run(f"[Изображение не найдено: {figure_id}]")
            idx += 1
            continue

        caption_match = CAPTION_RE.match(stripped)
        if caption_match:
            flush_paragraph_buffer(anchor, paragraph_buffer)
            caption = caption_match.group(1)
            para = insert_paragraph_before(anchor, cleanup_inline_markup(caption), style="Caption")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            idx += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph_buffer(anchor, paragraph_buffer)
            rows, idx = parse_markdown_table(lines, idx)
            insert_table_before(anchor, rows)
            continue

        if stripped.startswith("- "):
            flush_paragraph_buffer(anchor, paragraph_buffer)
            while idx < len(lines) and lines[idx].strip().startswith("- "):
                item = cleanup_inline_markup(lines[idx].strip()[2:])
                para = insert_paragraph_before(anchor, item, style="List Paragraph")
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                idx += 1
            continue

        if NUMBERED_RE.match(stripped):
            flush_paragraph_buffer(anchor, paragraph_buffer)
            while idx < len(lines) and NUMBERED_RE.match(lines[idx].strip()):
                item = cleanup_inline_markup(lines[idx].strip())
                para = insert_paragraph_before(anchor, item, style="List Paragraph")
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                idx += 1
            continue

        paragraph_buffer.append(stripped)
        idx += 1

    flush_paragraph_buffer(anchor, paragraph_buffer)
    set_update_fields_on_open(doc)
    doc.save(str(OUTPUT_DOCX))
    return OUTPUT_DOCX


if __name__ == "__main__":
    out = build_docx()
    print(out)
