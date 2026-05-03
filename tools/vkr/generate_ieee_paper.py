#!/usr/bin/env python3
"""
Generate MAIN 2026 conference paper using the official IEEE A4 template.
v5 – strict IEEE first-appearance citation order, figure renumbering,
     style softening, bibliography corrections.
"""

import os, copy, zipfile
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# ── paths ─────────────────────────────────────────────────────
BASE        = '/home/afetz/work/clean/NEWWAY/conference'
TEMPLATE_ORIG = os.path.join(BASE, 'conference-template-a4.docx')
TEMPLATE_TRANS = os.path.join(BASE, 'conference-template-a4-transitional.docx')
OUTPUT        = os.path.join(BASE, 'Fizulin_Romanov_MAIN2026_NRV2X_CoSim.docx')

FIG_DIR_SWEEP = '/home/afetz/work/clean/NEWWAY/analysis/mode2_loss/figures/sweep'
FIG_DIR_VKR   = '/home/afetz/work/clean/NEWWAY/archive/2026-05-03/vkr_manuscript/figures'
FIG1 = os.path.join(FIG_DIR_SWEEP, 'prr_vs_txPower.png')
FIG2 = os.path.join(FIG_DIR_SWEEP, 'reaction_delay_p90_vs_prr.png')
FIG3 = os.path.join(FIG_DIR_SWEEP, 'behavior_vs_prr_time_to_first_brake.png')
FIG4 = os.path.join(FIG_DIR_VKR,   'figure_2_1_architecture.png')

# ── Step 0: Convert OOXML Strict → Transitional if needed ────
STRICT_TO_TRANS = {
    'http://purl.oclc.org/ooxml/officeDocument/relationships/officeDocument':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/extendedProperties':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/customXml':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/styles':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/numbering':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/settings':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/webSettings':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/footnotes':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/endnotes':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/footer':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/header':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/hyperlink':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/fontTable':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/theme':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme',
    'http://purl.oclc.org/ooxml/officeDocument/relationships/image':
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
}

def convert_strict_to_transitional(src, dst):
    with zipfile.ZipFile(src, 'r') as zin, zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith('.xml') or item.filename.endswith('.rels'):
                text = data.decode('utf-8')
                for strict, trans in STRICT_TO_TRANS.items():
                    text = text.replace(strict, trans)
                text = text.replace('http://purl.oclc.org/ooxml/wordprocessingml/main',
                                    'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
                text = text.replace('http://purl.oclc.org/ooxml/drawingml/2006/main',
                                    'http://schemas.openxmlformats.org/drawingml/2006/main')
                text = text.replace('http://purl.oclc.org/ooxml/officeDocument/relationships',
                                    'http://schemas.openxmlformats.org/officeDocument/2006/relationships')
                text = text.replace('http://purl.oclc.org/ooxml/', 'http://schemas.openxmlformats.org/')
                data = text.encode('utf-8')
            zout.writestr(item, data)

convert_strict_to_transitional(TEMPLATE_ORIG, TEMPLATE_TRANS)

# ── Step 1: Open converted template ──────────────────────────
doc = Document(TEMPLATE_TRANS)

# ── Step 2: Remove ALL body content but preserve final sectPr ─
body = doc.element.body
ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
final_sectPr = body.find(f'{ns_w}sectPr')  # direct child sectPr at end

for child in list(body):
    if child is not final_sectPr:
        body.remove(child)

# ── Step 3: Remove footer content ────────────────────────────
for section in doc.sections:
    footer = section.footer
    if footer:
        for p in footer.paragraphs:
            for run in p.runs:
                run.text = ''
    # Unlink footer so it doesn't inherit
    section.footer.is_linked_to_previous = False
    for p in section.footer.paragraphs:
        for r in p.runs:
            r.text = ''

# ── Step 4: Set final section to 2-col with correct margins ──
if final_sectPr is not None:
    # Margins
    pgMar = final_sectPr.find(f'{ns_w}pgMar')
    if pgMar is not None:
        pgMar.set(qn('w:top'), '27pt')
        pgMar.set(qn('w:bottom'), '72pt')
        pgMar.set(qn('w:left'), '44.5pt')
        pgMar.set(qn('w:right'), '44.5pt')
    # 2 columns
    cols = final_sectPr.find(f'{ns_w}cols')
    if cols is not None:
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '240')
    else:
        final_sectPr.append(parse_xml(
            f'<w:cols {nsdecls("w")} w:num="2" w:space="240"/>'))
    # Remove any footer reference
    for ftr_ref in final_sectPr.findall(f'{ns_w}footerReference'):
        final_sectPr.remove(ftr_ref)

# ── helpers ───────────────────────────────────────────────────
def _add_p(text, style_name):
    return doc.add_paragraph(text, style=style_name)

def add_body(text):
    return _add_p(text, 'Body Text')

def add_ref(text):
    return _add_p(text, 'references')

def add_figure(image_path, caption, width=Cm(7.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    if os.path.exists(image_path):
        run.add_picture(image_path, width=width)
    else:
        run.add_text(f'[IMAGE NOT FOUND: {image_path}]')
    _add_p(caption, 'figure caption')

def add_table_ieee(headers, data, caption):
    _add_p(caption, 'table head')
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(data), cols=n_cols)
    table.style = 'Normal Table'
    # Borders
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tbl_pr)
    tbl_pr.append(parse_xml(borders_xml))
    # Header
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.size = Pt(8); r.font.bold = True; r.font.name = 'Times New Roman'
    # Data
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            c = table.rows[i+1].cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(8); r.font.name = 'Times New Roman'
    doc.add_paragraph()

def insert_section_break_continuous_2col():
    """Insert a continuous section break switching to 2 columns.
    This goes into the pPr of the LAST paragraph added so far."""
    # Create a new paragraph that carries the section break
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sectPr = parse_xml(
        f'<w:sectPr {nsdecls("w")}>'
        f'  <w:pgSz w:w="11907" w:h="16839" w:orient="portrait"/>'
        f'  <w:pgMar w:top="27pt" w:right="44.5pt" w:bottom="72pt" w:left="44.5pt"'
        f'           w:header="36pt" w:footer="36pt" w:gutter="0pt"/>'
        f'  <w:cols w:num="1"/>'
        f'  <w:type w:val="continuous"/>'
        f'</w:sectPr>'
    )
    pPr.append(sectPr)


# ============================================================
#  SECTION 1: FULL-WIDTH TITLE BLOCK  (1 column)
# ============================================================

# Title
_add_p(
    'Impact of 5G NR-V2X Mode 2 Message Losses on '
    'Cooperative Vehicle Behavior: A Reproducible '
    'Co-Simulation Study with Ray-Tracing-in-the-Loop',
    'paper title'
)

# Author 1
p = _add_p('', 'Author')
for text, sz, italic in [
    ('Andrey V. Fizulin', Pt(11), False),
    ('\nMoscow Institute of Electronics and Mathematics (MIEM)', Pt(10), True),
    ('\nHSE University, Moscow, Russia', Pt(10), True),
    ('\navfizulin@edu.hse.ru', Pt(10), False),
]:
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = sz
    r.font.italic = italic

# Author 2
p = _add_p('', 'Author')
for text, sz, italic in [
    ('Aleksandr Yu. Romanov', Pt(11), False),
    ('\nMoscow Institute of Electronics and Mathematics (MIEM)', Pt(10), True),
    ('\nHSE University, Moscow, Russia', Pt(10), True),
    ('\nayromanov@hse.ru', Pt(10), False),
]:
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = sz
    r.font.italic = italic

# ── section break: continuous, switch from 1-col to 2-col ────
insert_section_break_continuous_2col()

# ============================================================
#  SECTION 2: TWO-COLUMN BODY  (abstract → references)
#
#  Citation numbering: strict IEEE first-appearance order [1]–[24].
#  Figure numbering: strict order of first mention in text.
#    Fig. 1 = architecture (III.A)
#    Fig. 2 = PRR vs txPower (V.A)
#    Fig. 3 = reaction delay (V.B)
#    Fig. 4 = braking time (V.C)
# ============================================================

# Abstract
_add_p(
    "Abstract\u2014Connected and automated vehicles rely on timely V2X message exchange "
    "for cooperative safety functions, yet the quantitative link between 5G NR sidelink "
    "Mode\u00a02 radio-layer degradation and application-level behavioral outcomes remains "
    "insufficiently explored. We present a reproducible co-simulation framework coupling "
    "SUMO, ns-3 with NR-V2X, and NVIDIA Sionna\u00a0RT ray-tracing, augmented with an "
    "ID-aware packet tracing mechanism and an automated causal audit pipeline. Two "
    "deterministic scenarios are evaluated: an Emergency Vehicle Alert with parametric "
    "transmit-power variation, and a lane-change-upon-leader-stop scenario with "
    "per-vehicle link-quality degradation. Results show that reducing the packet "
    "reception ratio from 0.94 to 0.08 increases the 90th-percentile cooperative "
    "reaction delay from 4.0\u00a0s to 18.1\u00a0s, delays first braking by up to 5\u00a0s, "
    "and, in the lane-change scenario, the causal audit confirms a complete chain from "
    "packet loss through delayed decision to collision. These findings quantify the "
    "sensitivity of cooperative driving functions to sidelink quality and support "
    "a digital-twin methodology for tracing communication failures through to "
    "safety-critical traffic outcomes.",
    'Abstract'
)

# Keywords
_add_p(
    "Keywords\u2014V2X, NR-V2X Mode 2, co-simulation, digital twin, packet reception ratio, "
    "cooperative driving, ray-tracing, SUMO, ns-3, Sionna RT, causal audit.",
    'Keywords'
)

# ── I. INTRODUCTION ──
_add_p('Introduction', 'Heading 1')

add_body(
    "Connected and automated vehicles (CAVs) introduce fundamentally new requirements "
    "for the quality of telecommunications infrastructure. Cooperative perception, hazard "
    "warning, and maneuver coordination depend on timely exchange of V2X (Vehicle-to-Everything) "
    "messages, and any disruptions in the communication channel directly affect driving "
    "safety [1], [2], [3]. In the European ITS architecture, Cooperative Awareness Messages "
    "(CAMs) provide periodic kinematic state updates [4], while Decentralized Environmental "
    "Notification Messages (DENMs) deliver event-driven alerts [5]."
)

add_body(
    "Traditional approaches to V2X quality assessment rely on aggregate network metrics\u2014"
    "mean packet reception ratio (PRR) and average latency\u2014without establishing a formal "
    "link between network events and observed traffic outcomes [6], [7]. Consequently, the "
    "assessment of how message losses affect cooperative driving safety remains "
    "fragmented: network metrics are analyzed in isolation from behavioral metrics, and "
    "the causal chain \u2018message loss \u2192 changed agent decision \u2192 traffic outcome\u2019 "
    "is neither formalized nor experimentally verified."
)

add_body(
    "For 5G NR-V2X sidelink in autonomous resource allocation mode (Mode\u00a02), the loss "
    "structure is particularly complex. Beyond external radio interference, losses arise from "
    "distributed access collisions, hidden nodes, resource selection window parameters, and "
    "application traffic load [8]. The 3GPP evaluation methodology TR\u00a037.885 defines "
    "PRR as the baseline reliability metric for V2X scenarios [9], yet equal PRR values "
    "may correspond to fundamentally different internal loss structures\u2014rendering PRR "
    "necessary but insufficient for explaining application-level behavioral transitions [7], [8]."
)

add_body(
    "Existing simulation environments for CAVs often decouple communication quality "
    "assessment from behavioral modeling. While frameworks like Veins [10], Eclipse "
    "MOSAIC [11], and ms-van3t [12] provide varying degrees of integration, the link "
    "between channel degradation and safety-critical traffic outcomes through intermediate "
    "metrics (packet inter-reception interval, age of information) is rarely traced end-to-end. "
    "Ray-tracing-in-the-loop approaches, such as VaN3Twin [13], strengthen the geometric "
    "fidelity of channel modeling but do not inherently provide the causal traceability "
    "from individual packet losses to behavioral decisions."
)

add_body(
    "Bridging this gap requires a digital-twin co-simulation methodology that embeds "
    "data-driven causal analysis directly into the evaluation loop of intelligent networked "
    "cyber-physical systems (CPS). Such a methodology reflects the growing consensus "
    "that application-aware network assessment\u2014where the network is evaluated not by "
    "aggregate KPIs but by its downstream impact on autonomous agent behavior\u2014is "
    "important for the design and assessment of cooperative driving functions [2]."
)

add_body(
    "This paper makes the following contributions: (1)\u00a0a reproducible co-simulation "
    "framework integrating SUMO, ns-3 with NR-V2X Mode\u00a02, and NVIDIA Sionna\u00a0RT "
    "ray-tracing, with an ID-aware packet tracing mechanism that achieves strict match "
    "ratio\u00a0=\u00a01.0; (2)\u00a0an automated causal audit pipeline linking packet drops to "
    "application-level decisions and traffic outcomes; (3)\u00a0two deterministic experimental "
    "scenarios quantifying the nonlinear relationship between PRR degradation and "
    "cooperative reaction delay, braking time, and collision occurrence; and (4)\u00a0an "
    "experimental demonstration that the 90th-percentile reaction delay increases from "
    "3.98\u00a0s to 18.09\u00a0s as PRR degrades from 0.94 to 0.077."
)

# ── II. RELATED WORK ──
_add_p('Related Work', 'Heading 1')

add_body(
    "The influence of communication quality on CAV behavior has been studied along "
    "several axes. Viriyasitavat et al. [6] survey channel and propagation models for "
    "vehicular communications, emphasizing the strong dependence of reliability on the "
    "environment and chosen model. Boban et al. [14] propose GEMV\u00b2, a geometry-based "
    "V2V channel model suitable for large-scale simulation, differentiating line-of-sight, "
    "vehicle-obstructed [15], and building-obstructed conditions."
)

add_body(
    "For NR-V2X sidelink, Harounabadi et al. [8] review Mode\u00a02 resource allocation "
    "in Release\u00a016 and beyond, highlighting distributed scheduling challenges. "
    "Romanov and Stepanyants [2] survey integrated simulation environments for CAVs, "
    "identifying requirements for reproducibility, tool composition, and architecture. "
    "The ms-van3t framework [12] provides a multi-stack V2X validation environment "
    "atop ns-3, while VaN3Twin [13] extends it with Sionna\u00a0RT ray-tracing-in-the-loop."
)

add_body(
    "Surrogate safety assessment using TTC and PET is well-established [16], and the "
    "role of Age of Information (AoI) as a bridge metric between network events and "
    "controller state has been formalized in [7], [17]. However, prior works typically "
    "evaluate network metrics and behavioral metrics separately. To the best of our "
    "knowledge, no prior study provides end-to-end causal tracing from individual "
    "5G NR-V2X Mode\u00a02 packet drops through application-level decisions to collision "
    "outcomes within a reproducible digital-twin framework with ray-tracing."
)

# ── III. CO-SIMULATION FRAMEWORK ──
_add_p('Co-Simulation Framework', 'Heading 1')
_add_p('Architecture Overview', 'Heading 2')

add_body(
    "The proposed framework implements a bidirectional cyber-physical simulation loop "
    "comprising four layers: (1)\u00a0the application layer (emergencyVehicleAlert), "
    "(2)\u00a0the ETSI Facilities Layer (CA Basic Service, DEN Basic Service, "
    "CP Basic Service [18]), "
    "(3)\u00a0the network/transport layer (BTP, GeoNetworking, DCC), and (4)\u00a0the "
    "channel/physical layer supporting IEEE\u00a0802.11p, LTE-V2X Mode\u00a04, and NR-V2X "
    "Mode\u00a02 [19] with optional Sionna\u00a0RT ray-tracing. The traffic simulator SUMO "
    "communicates with ns-3 via the TraCI interface [20], synchronizing vehicle positions, "
    "speeds, and control actions at each simulation step."
)

add_figure(FIG4,
    'Fig. 1. Architecture of the co-simulation framework. '
    'Arrows denote data flow between SUMO (traffic), ns-3 (network), '
    'Sionna RT (channel), and the causal audit pipeline.',
    width=Cm(8.0))

add_body(
    "Sionna\u00a0RT [21] is used as the ray-tracing engine. In our implementation, it is "
    "deployed as an external GPU service that accepts UDP requests for "
    "channel-parameter computation given 3D transmitter/receiver positions. At each "
    "ns-3 step, the module submits vehicle coordinates; Sionna performs ray-tracing "
    "accounting for scene geometry and returns channel parameters (path loss, Doppler "
    "shift). This design follows the VaN3Twin ray-tracing-in-the-loop paradigm [13] "
    "while extending it with per-vehicle link-quality profiles and causal tracing. "
    "The resulting digital-twin loop enables application-aware network assessment: "
    "the same simulation instance that computes PHY-layer metrics also drives the "
    "traffic agents whose safety outcomes are the ultimate evaluation criterion."
)

_add_p('Per-Vehicle Link-Quality Profiles', 'Heading 2')

add_body(
    "A key extension is the per-vehicle-prr-profile mechanism, which assigns each vehicle "
    "individual radio parameters: PHY-layer CAM drop probability, equivalent transmit "
    "power, and target PRR. Profiles are specified as command-line parameters in the "
    "format veh3:0.05:23:0.95,veh4:0.923:\u221220:0.077. This enables controlled, "
    "reproducible degradation of specific links while maintaining identical traffic "
    "scenarios, initial conditions, and random number generator states."
)

_add_p('ID-Aware Packet Tracing', 'Heading 2')

add_body(
    "Each transmitted CAM is assigned a unique packet identifier (pkt_uid) preserved "
    "across all lifecycle stages: transmission (TX), successful reception (RX), PHY-layer "
    "drop (DROP_PHY), and application-level decision (CTRL). The tracing system "
    "generates per-vehicle log files: MSG.csv (message events), CTRL.csv (control "
    "decisions), and PROFILE.csv (link-quality configuration). A post-processing script "
    "correlates DROP events with DECISION events by pkt_uid, constructing a "
    "timeline that maps each lost packet to its corresponding application-level "
    "consequence. The achieved strict match ratio\u00a0=\u00a01.0 confirms that every "
    "packet loss has a traceable application-level outcome."
)

_add_p('Automated Causal Audit', 'Heading 2')

add_body(
    "The causal audit pipeline reconstructs the chain \u2018packet loss \u2192 no-action "
    "decision \u2192 collision\u2019 for each crash event. The audit script identifies collision "
    "participants from SUMO\u2019s collision.xml, analyzes the event window preceding the "
    "collision timestamp, counts drop events and no-action decisions within that window, "
    "and classifies the causal strength as strong_no_action_only, mixed, or weak. "
    "This automated, data-driven procedure replaces manual log inspection and provides a "
    "structured, reproducible evidence chain for attributing safety outcomes to "
    "communication failures."
)

# ── IV. EXPERIMENTAL DESIGN ──
_add_p('Experimental Design', 'Heading 1')
_add_p('Scenario 1: Emergency Vehicle Alert with Parametric TX Power Sweep', 'Heading 2')

add_body(
    "The first scenario is based on the v2v-emergencyVehicleAlert-nrv2x example from "
    "the open-source ms-van3t framework [12]. In NR-V2X Mode\u00a02 sidelink "
    "configuration, vehicles exchange CAMs [4] using autonomous resource allocation "
    "[8], [22]. A parametric sweep varies only the transmit power (txPower) across "
    "iterations while fixing trajectories, parameters, and random number generator seeds. "
    "The communication metric is PRR, defined as the fraction of successful receptions "
    "among potential receivers [9]. The reaction metric is the time to first CAM reception "
    "from the emergency vehicle (stationId\u00a0=\u00a02) for each participant, from which median "
    "and 90th-percentile statistics are computed. Additionally, SUMO data provide the "
    "median time to first braking."
)

_add_p('Scenario 2: Lane-Change upon Leader Stop with Causal Audit', 'Heading 2')

add_body(
    "The second scenario models a sequence of vehicles in the right lane, where the "
    "leading vehicle stops, becoming a stationary obstacle. The experimental hypothesis "
    "states that if following vehicles receive timely V2X warnings, they perform lane "
    "changes and maintain safe traffic flow; if one vehicle\u2019s link quality degrades sharply, "
    "its decision is delayed and the conflict results in a collision."
)

add_body(
    "Two contrasting regimes are defined. In the safe regime, all vehicles maintain high "
    "link quality: veh3 (equiv_tx_power\u00a0=\u00a023\u00a0dBm, PRR\u00a0=\u00a00.945), "
    "veh4 (23\u00a0dBm, PRR\u00a0=\u00a00.944), veh5 (0\u00a0dBm, PRR\u00a0=\u00a00.722). "
    "In the crash regime, degradation is introduced only for veh4: "
    "equiv_tx_power\u00a0=\u00a0\u221220\u00a0dBm, PRR\u00a0=\u00a00.071, while veh3 "
    "and veh5 retain identical parameters. The traffic scenario, incident time "
    "(t\u00a0=\u00a06.00\u00a0s), and fleet composition remain constant; only the critical "
    "vehicle\u2019s link quality changes."
)

_add_p('Reproducibility Provisions', 'Heading 2')

add_body(
    "All experiments follow a strict reproducibility protocol [2]: road geometry and "
    "route files are versioned; SUMO microscopic model parameters are fixed; network "
    "parameters, link-quality profiles, Sionna\u00a0RT settings, and random number generator "
    "states are locked. For each scenario, not only final plots but also raw CSV logs are "
    "produced as auditable artifacts. The entire experiment is executable via a single "
    "shell command."
)

# ── V. RESULTS ──
_add_p('Results', 'Heading 1')

add_body(
    "The results presented below are drawn from a deterministic proof-of-concept study "
    "rather than a full statistical campaign. Each experimental point corresponds to a "
    "single controlled run with fixed seeds, enabling exact causal attribution at the "
    "cost of statistical generalizability. We consider this an appropriate trade-off for "
    "demonstrating the methodology and establishing the causal chain; multi-seed "
    "campaigns with confidence intervals are deferred to future work (Section\u00a0VII)."
)

_add_p('PRR vs. Transmit Power', 'Heading 2')

add_body(
    "Fig.\u00a02 shows the monotonic increase of PRR with txPower in NR-V2X Mode\u00a02. "
    "At txPower\u00a0=\u00a0\u221220\u00a0dBm, PRR drops to approximately 0.077, while at "
    "txPower\u00a0=\u00a023\u00a0dBm it reaches 0.94. The relationship is nonlinear, with "
    "a sharp transition in the range [\u221210,\u00a00]\u00a0dBm."
)

add_figure(FIG1,
    'Fig. 2. PRR vs. txPower (NR-V2X Mode 2). Monotonic increase with '
    'sharp transition around \u221210 to 0 dBm.',
    width=Cm(7.5))

_add_p('Cooperative Reaction Delay', 'Heading 2')

add_body(
    "Fig.\u00a03 presents the 90th-percentile cooperative reaction delay as a function of PRR. "
    "At PRR\u00a0\u2248\u00a00.94, the 90th-percentile delay is 3.98\u00a0s; at "
    "PRR\u00a0\u2248\u00a00.077, it rises to 18.09\u00a0s\u2014a factor of 4.5\u00d7. The "
    "relationship is strongly nonlinear: moderate PRR degradation (0.94\u00a0\u2192\u00a00.6) "
    "increases the delay by approximately 2\u00a0s, whereas degradation below 0.3 causes "
    "the delay to escalate sharply. This nonlinearity implies that system designers cannot "
    "rely on linear interpolation of delay bounds from PRR measurements."
)

add_figure(FIG2,
    'Fig. 3. 90th-percentile cooperative reaction delay vs. PRR. Nonlinear '
    'escalation below PRR = 0.3.',
    width=Cm(7.5))

_add_p('Time to First Braking', 'Heading 2')

add_body(
    "Fig.\u00a04 shows the median time to first braking as a function of PRR. The trend "
    "is consistent with the reaction delay: at high PRR, braking initiates near t\u00a0\u2248\u00a06\u00a0s "
    "(shortly after the incident), whereas at PRR\u00a0\u2248\u00a00.077, the median braking "
    "onset shifts to t\u00a0\u2248\u00a011\u00a0s. The 5-second delay in braking onset represents "
    "a substantial safety degradation, particularly at highway speeds where each second "
    "of delayed braking corresponds to approximately 25\u201333\u00a0m of additional travel."
)

add_figure(FIG3,
    'Fig. 4. Median time to first braking vs. PRR. Delayed braking at low PRR '
    'extends stopping distance.',
    width=Cm(7.5))

_add_p('Safe vs. Crash Regime: Causal Analysis', 'Heading 2')

add_body(
    "In Scenario\u00a02, the safe regime produces an orderly cascade of lane changes: "
    "veh3 at t\u00a0=\u00a07.14\u00a0s, veh4 at t\u00a0=\u00a09.13\u00a0s, veh5 at "
    "t\u00a0=\u00a012.13\u00a0s, with zero collisions in collision.xml. In the crash regime, "
    "veh3 still changes lane successfully (high PRR), but veh4 (PRR\u00a0=\u00a00.071) "
    "fails to receive the emergency CAM in time. The causal audit confirms a "
    "strong_no_action_only classification: every dropped CAM maps to a "
    "no-action decision on the application layer, and the accumulation of "
    "no-action decisions directly precedes the collision event."
)

add_body(
    "Table\u00a0I summarizes the key timing differences between the two regimes. "
    "The critical observation is that the traffic scenario is identical in both cases\u2014"
    "the only variable is veh4\u2019s equivalent transmit power, which controls its PRR. "
    "This isolation of the communication variable, combined with the pkt_uid-based "
    "causal audit, provides strong evidence that the collision outcome is causally "
    "attributable to V2X link degradation."
)

add_table_ieee(
    ['Metric', 'Safe Regime', 'Crash Regime'],
    [
        ['veh3 first lane change', '7.14 s', '7.14 s'],
        ['veh4 first lane change', '9.13 s', 'No lane change'],
        ['veh5 first lane change', '12.13 s', '12.13 s'],
        ['Collision', 'None', 'veh4 rear-ends veh2'],
        ['Causal classification', 'N/A', 'strong_no_action_only'],
    ],
    'TABLE I. Timing comparison: safe vs. crash regime (Scenario 2)',
)

_add_p('Cross-Technology Comparison', 'Heading 2')

add_body(
    "Table\u00a0II provides a comparative snapshot of IEEE\u00a0802.11p and NR-V2X Mode\u00a02 "
    "under identical co-existence conditions (10 nodes, standard parameters) [23]. NR-V2X "
    "achieves a higher mean PRR (0.962 vs. 0.920) but at the cost of substantially higher "
    "latency (12.3\u00a0ms vs. 0.49\u00a0ms). For safety-critical cooperative functions, this "
    "latency\u2013reliability trade-off must be evaluated in the context of application-level "
    "timing budgets rather than aggregate network metrics alone. "
    "The values in Table\u00a0II are configuration-specific and reflect the considered "
    "co-existence setup and parameterization rather than a universal ranking of the "
    "two radio technologies; different parameterizations or traffic loads may yield "
    "different results."
)

add_table_ieee(
    ['Metric', 'IEEE 802.11p', 'NR-V2X Mode 2'],
    [
        ['Mean PRR', '0.920', '0.962'],
        ['PRR range', '0.789 \u2013 1.000', '0.826 \u2013 1.000'],
        ['Mean latency', '0.49 ms', '12.3 ms'],
        ['Latency range', '0.44 \u2013 0.64 ms', '7.7 \u2013 16.8 ms'],
    ],
    'TABLE II. Comparison of radio technologies (10 nodes, co-existence)',
)

# ── VI. DISCUSSION ──
_add_p('Discussion', 'Heading 1')

add_body(
    "The experimental results reveal several insights of practical importance. "
    "First, the nonlinear sensitivity of cooperative reaction delay to PRR degradation "
    "implies that performance guarantees based on mean PRR are misleading. The "
    "90th-percentile delay\u2014representing the worst-case participants\u2014escalates "
    "much faster than the mean, particularly below PRR\u00a0=\u00a00.3. System designers "
    "must therefore specify communication requirements in terms of tail percentile "
    "bounds on reaction delay, not average PRR."
)

add_body(
    "Second, the causal audit methodology provides a principled mechanism for "
    "attributing safety outcomes to communication failures. The strict match "
    "ratio\u00a0=\u00a01.0 between DROP_PHY events and application-level decisions "
    "indicates that, in the considered scenarios, every packet loss has a traceable "
    "application-level outcome. The classification of causal "
    "strength (strong, mixed, weak) enables automated prioritization of failure "
    "modes in large-scale simulation campaigns."
)

add_body(
    "Third, the per-vehicle link-quality profile mechanism enables fine-grained "
    "controlled experiments in which only the communication variable changes while "
    "the traffic scenario remains identical. This is important for isolating causal "
    "effects and distinguishes our approach from studies that vary traffic and "
    "communication parameters simultaneously."
)

add_body(
    "Fourth, the integration of Sionna\u00a0RT ray-tracing provides "
    "geometry-sensitive channel modeling, capturing effects of buildings, vehicle "
    "obstructions, and scene geometry that stochastic models may average out. "
    "For intersection scenarios where line-of-sight transitions are abrupt, this "
    "level of fidelity is important for realistic assessment of V2X reliability."
)

add_body(
    "From a broader perspective, the framework instantiates a digital-twin approach to "
    "intelligent networked CPS evaluation: the co-simulation loop reproduces the full "
    "cyber-physical interaction between communication, perception, decision-making, and "
    "vehicle dynamics, while the automated causal audit provides a data-driven analytical "
    "layer that converts raw simulation logs into structured safety evidence. This "
    "combination offers a basis for systematic assessment of cooperative driving systems, "
    "where reproducible, end-to-end causal traceability is highly desirable."
)

# ── VII. LIMITATIONS AND FUTURE WORK ──
_add_p('Limitations and Future Work', 'Heading 1')

add_body(
    "Several limitations apply. The presented results are derived from single-seed "
    "deterministic runs; while this design maximizes causal clarity, it does not yield "
    "confidence intervals or distributional statistics. A multi-seed statistical campaign "
    "covering varying traffic densities, radio parameters, and random access contention "
    "levels is required to generalize the quantitative findings."
)

add_body(
    "The vehicle control logic is not closed-loop with respect to network metrics: "
    "trajectory modifications and maneuver stability under varying feedback quality "
    "are not examined. The analysis focuses on the reaction to the first "
    "successfully received CAM and derived percentile statistics."
)

add_body(
    "The per-vehicle-prr-profile mechanism, while enabling controlled degradation, "
    "applies a simplified loss model (independent drops). In NR-V2X Mode\u00a02, "
    "burst losses due to persistent half-duplex collisions are structurally different "
    "from independent losses [8]. Future work will integrate burst-aware loss "
    "patterns derived from MAC-layer traces."
)

add_body(
    "Planned extensions include: (1)\u00a0expanding the causal audit to incorporate "
    "PIR and AoI as intermediate causal variables; (2)\u00a0evaluating additional cooperative "
    "functions (platooning [24], cooperative intersection management); (3)\u00a0conducting "
    "multi-seed statistical campaigns with confidence intervals; and (4)\u00a0validating "
    "simulation results against field measurements from V2X testbeds."
)

# ── VIII. CONCLUSION ──
_add_p('Conclusion', 'Heading 1')

add_body(
    "We presented a reproducible co-simulation framework for quantifying the impact "
    "of 5G NR-V2X Mode\u00a02 message losses on cooperative vehicle behavior. The "
    "framework couples SUMO, ns-3, and Sionna\u00a0RT ray-tracing with an ID-aware "
    "packet tracing mechanism achieving strict match ratio\u00a0=\u00a01.0 and an automated "
    "causal audit pipeline."
)

add_body(
    "Experiments on two deterministic scenarios show that: (1)\u00a0PRR degradation "
    "from 0.94 to 0.077 increases the 90th-percentile cooperative reaction delay from "
    "3.98\u00a0s to 18.09\u00a0s, exhibiting strongly nonlinear sensitivity; (2)\u00a0median braking "
    "onset is delayed by up to 5\u00a0s under severe link degradation; (3)\u00a0the causal audit "
    "confirms a \u2018packet loss \u2192 no-action \u2192 collision\u2019 chain with strong causal "
    "classification in the crash regime; (4)\u00a0NR-V2X Mode\u00a02 offers higher PRR than "
    "IEEE\u00a0802.11p but at higher latency, requiring application-specific evaluation."
)

add_body(
    "These results indicate that cooperative safety functions are sensitive to "
    "communication channel quality in ways that aggregate metrics can obscure, and that "
    "reproducible co-simulation with end-to-end causal tracing provides a robust basis "
    "for reliable V2X system assessment. The framework source code, scenario "
    "configurations, and analysis scripts will be made available upon publication."
)

# ── ACKNOWLEDGMENT (unnumbered — Heading 5) ──
_add_p('Acknowledgment', 'Heading 5')

add_body(
    "This work was conducted at the CAD Systems Laboratory, Department of Computer "
    "Engineering, Moscow Institute of Electronics and Mathematics (MIEM), HSE University."
)

# ── REFERENCES (unnumbered — Heading 5) ──
_add_p('References', 'Heading 5')

# References ordered strictly by first appearance in text.
# The "references" style provides automatic [n] numbering via numbering.xml.
refs = [
    # [1] SAE J3016
    'SAE International, \u201cTaxonomy and Definitions for Terms Related to Driving '
    'Automation Systems for On-Road Motor Vehicles,\u201d SAE J3016, Rev. Apr. 2021, '
    'doi: 10.4271/J3016_202104.',

    # [2] Stepanyants & Romanov
    'V. G. Stepanyants and A. Yu. Romanov, \u201cA survey of integrated simulation '
    'environments for connected automated vehicles: Requirements, tools, and '
    'architecture,\u201d IEEE Intell. Transp. Syst. Mag., vol. 16, no. 2, pp. 6\u201322, '
    '2024, doi: 10.1109/MITS.2023.3335126.',

    # [3] NHTSA V2V
    'NHTSA, \u201cVehicle-to-Vehicle Communication: Readiness of V2V Technology for '
    'Application,\u201d DOT HS 812 014, 2014.',

    # [4] ETSI CAM
    'ETSI EN 302 637-2 V1.4.1, \u201cIntelligent Transport Systems (ITS); Vehicular '
    'Communications; Basic Set of Applications; Part 2: Specification of Cooperative '
    'Awareness Basic Service,\u201d 2019.',

    # [5] ETSI DENM
    'ETSI EN 302 637-3 V1.3.1, \u201cIntelligent Transport Systems (ITS); Vehicular '
    'Communications; Basic Set of Applications; Part 3: Specification of Decentralized '
    'Environmental Notification Basic Service,\u201d 2019.',

    # [6] Viriyasitavat
    'W. Viriyasitavat, M. Boban, H.-M. Tsai, and A. Vasilakos, \u201cVehicular '
    'communications: Survey and challenges of channel and propagation models,\u201d '
    'IEEE Veh. Technol. Mag., vol. 10, no. 2, pp. 55\u201366, Jun. 2015, '
    'doi: 10.1109/MVT.2015.2410341.',

    # [7] Abd-Elmagid (AoI)
    'M. A. Abd-Elmagid, N. Pappas, and H. S. Dhillon, \u201cOn the role of age of '
    'information in the Internet of Things,\u201d IEEE Commun. Mag., vol. 57, no. 12, '
    'pp. 72\u201377, Dec. 2019, doi: 10.1109/MCOM.001.1900041.',

    # [8] Harounabadi
    'M. Harounabadi et al., \u201cV2X in 3GPP standardization: NR sidelink in Rel-16 '
    'and beyond,\u201d IEEE Commun. Standards Mag., vol. 5, no. 1, pp. 12\u201321, 2021, '
    'doi: 10.1109/MCOMSTD.001.2000070.',

    # [9] 3GPP TR 37.885
    '3GPP TR 37.885, \u201cStudy on evaluation methodology of new V2X use cases '
    'for LTE and NR,\u201d v15.3.0, 2019.',

    # [10] Sommer / Veins
    'C. Sommer, R. German, and F. Dressler, \u201cBidirectionally Coupled Network and '
    'Road Traffic Simulation for Improved IVC Analysis,\u201d IEEE Trans. Mobile Comput., '
    'vol. 10, no. 1, pp. 3\u201315, Jan. 2011, doi: 10.1109/TMC.2010.133.',

    # [11] Eclipse MOSAIC
    'Eclipse MOSAIC, \u201cMulti-Domain and Multi-Scale Simulation for Connected and '
    'Automated Mobility,\u201d [Online]. Available: https://eclipse.dev/mosaic/',

    # [12] ms-van3t
    'DriveX-devs, \u201cms-van3t: Multi-stack vehicular network simulation framework,\u201d '
    '[Online]. Available: https://github.com/DriveX-devs/ms-van3t',

    # [13] VaN3Twin
    'DriveX-devs, \u201cVaN3Twin: Vehicular digital twin framework with ray-tracing-in-'
    'the-loop,\u201d [Online]. Available: https://github.com/DriveX-devs/VaN3Twin',

    # [14] Boban GEMV\u00b2
    'M. Boban, J. Barros, and O. K. Tonguz, \u201cGeometry-based vehicle-to-vehicle '
    'channel modeling for large-scale simulation,\u201d IEEE Trans. Veh. Technol., vol. 63, '
    'no. 9, pp. 4146\u20134164, Nov. 2014, doi: 10.1109/TVT.2014.2317803.',

    # [15] Boban obstacles
    'M. Boban, T. T. V. Vinhoza, M. Ferreira, J. Barros, and O. K. Tonguz, '
    '\u201cImpact of vehicles as obstacles in vehicular ad hoc networks,\u201d IEEE J. Sel. '
    'Areas Commun., vol. 29, no. 1, pp. 15\u201328, Jan. 2011, '
    'doi: 10.1109/JSAC.2011.110103.',

    # [16] FHWA surrogate safety
    'FHWA, \u201cSurrogate Safety Assessment Model and Validation: Final Report,\u201d '
    'FHWA-HRT-08-051, 2008.',

    # [17] Sun (AoI)
    'Y. Sun, E. Uysal-Biyikoglu, R. D. Yates, C. E. Koksal, and N. B. Shroff, '
    '\u201cUpdate or wait: How to keep your data fresh,\u201d IEEE Trans. Inf. Theory, vol. 63, '
    'no. 11, pp. 7492\u20137508, Nov. 2017, doi: 10.1109/TIT.2017.2735804.',

    # [18] ETSI CPS
    'ETSI TS 103 324 V2.1.1, \u201cIntelligent Transport Systems (ITS); Vehicular '
    'Communications; Collective Perception Service; Release 2,\u201d 2023.',

    # [19] 3GPP TS 23.287
    '3GPP TS 23.287 V17.5.0, \u201cArchitecture enhancements for 5G System (5GS) to support '
    'Vehicle-to-Everything (V2X) services,\u201d Sep. 2022.',

    # [20] SUMO / Krajzewicz
    'D. Krajzewicz, J. Erdmann, M. Behrisch, and L. Bieker, \u201cRecent development '
    'and applications of SUMO \u2013 Simulation of Urban Mobility,\u201d Int. J. Adv. Syst. '
    'Meas., vol. 5, no. 3\u20134, pp. 128\u2013138, 2012.',

    # [21] NVIDIA Sionna
    'NVIDIA, \u201cSionna: An Open-Source Library for Next-Generation Physical Layer '
    'Research,\u201d [Online]. Available: https://developer.nvidia.com/sionna',

    # [22] 3GPP TS 38.321
    '3GPP TS 38.321, \u201cNR; Medium Access Control (MAC) protocol specification,\u201d '
    'Rel. 16, v16.7.0, 2021.',

    # [23] Molina-Masegosa
    'A. Molina-Masegosa, J. Gozalvez, and M. Sepulcre, \u201cComparison of IEEE '
    '802.11p and LTE-V2X: An evaluation with periodic and aperiodic messages of '
    'constant and variable size,\u201d IEEE Access, vol. 8, pp. 121526\u2013121548, 2020, '
    'doi: 10.1109/ACCESS.2020.3007115.',

    # [24] Heinovski / PlaFoSim
    'J. Heinovski, D. S. Buse, and F. Dressler, \u201cScalable Simulation of Platoon '
    'formation maneuvers with PlaFoSim,\u201d in Proc. IEEE Veh. Netw. Conf. (VNC), '
    'Poster Session, 2021, pp. 137\u2013138, doi: 10.1109/VNC52810.2021.9644678.',
]

for ref in refs:
    add_ref(ref)

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f'Paper saved to: {OUTPUT}')
print(f'File size: {os.path.getsize(OUTPUT):,} bytes')
